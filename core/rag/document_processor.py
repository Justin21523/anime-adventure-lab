# core/rag/document_processor.py

import os
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass
import hashlib
import mimetypes
from datetime import datetime
import re

# Document parsers
try:
    import docx
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import openpyxl
    from openpyxl import load_workbook

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

import opencc

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Processed document with metadata"""

    doc_id: str
    content: str
    metadata: Dict[str, Any]
    file_type: str
    file_size: int
    processed_at: datetime


class DocumentProcessor:
    """Document processor for various file formats"""

    def __init__(self):
        # Initialize text processors
        try:
            self.cc_t2s = opencc.OpenCC("t2s")  # Traditional to Simplified
            self.cc_s2t = opencc.OpenCC("s2t")  # Simplified to Traditional
        except:
            logger.warning("OpenCC not available")
            self.cc_t2s = None
            self.cc_s2t = None

    def process_file(
        self,
        file_path: Union[str, Path],
        doc_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """Process a file and extract text content"""

        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Generate doc_id if not provided
        if doc_id is None:
            doc_id = self._generate_doc_id(file_path)

        # Detect file type
        file_type = self._detect_file_type(file_path)
        file_size = file_path.stat().st_size

        # Initialize metadata
        if metadata is None:
            metadata = {}

        metadata.update(
            {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_type": file_type,
                "file_size": file_size,
                "processed_at": datetime.now().isoformat(),
            }
        )

        # Extract content based on file type
        content = self._extract_content(file_path, file_type)

        # Clean and normalize content
        content = self._clean_content(content)

        return ProcessedDocument(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_type=file_type,
            file_size=file_size,
            processed_at=datetime.now(),
        )

    def process_text(
        self,
        text: str,
        doc_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """Process raw text content"""

        if doc_id is None:
            doc_id = self._generate_text_id(text)

        if metadata is None:
            metadata = {}

        metadata.update(
            {
                "content_type": "raw_text",
                "char_count": len(text),
                "processed_at": datetime.now().isoformat(),
            }
        )

        # Clean and normalize content
        content = self._clean_content(text)

        return ProcessedDocument(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_type="text",
            file_size=len(text.encode("utf-8")),
            processed_at=datetime.now(),
        )

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID based on file"""
        # Use file path and modification time for unique ID
        file_stat = file_path.stat()
        content = f"{file_path.name}_{file_stat.st_size}_{file_stat.st_mtime}"
        return hashlib.md5(content.encode()).hexdigest()[:12]

    def _generate_text_id(self, text: str) -> str:
        """Generate unique ID for text content"""
        content_hash = hashlib.md5(text.encode("utf-8")).hexdigest()
        return f"text_{content_hash[:12]}"

    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type from extension and MIME type"""
        extension = file_path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(file_path))

        # Map common extensions to types
        type_mapping = {
            ".txt": "text",
            ".md": "markdown",
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".xlsx": "excel",
            ".xls": "excel",
            ".csv": "csv",
            ".json": "json",
            ".xml": "xml",
            ".html": "html",
            ".htm": "html",
        }

        return type_mapping.get(extension, "unknown")

    def _extract_content(self, file_path: Path, file_type: str) -> str:
        """Extract text content from file based on type"""

        try:
            if file_type == "text" or file_type == "markdown":
                return self._extract_text_file(file_path)

            elif file_type == "pdf":
                return self._extract_pdf(file_path)

            elif file_type == "docx":
                return self._extract_docx(file_path)

            elif file_type == "excel":
                return self._extract_excel(file_path)

            elif file_type == "csv":
                return self._extract_csv(file_path)

            elif file_type == "json":
                return self._extract_json(file_path)

            elif file_type == "html":
                return self._extract_html(file_path)

            else:
                # Try to read as plain text
                logger.warning(f"Unknown file type {file_type}, trying as text")
                return self._extract_text_file(file_path)

        except Exception as e:
            logger.error(f"Content extraction failed for {file_path}: {e}")
            return f"[Content extraction failed: {str(e)}]"

    def _extract_text_file(self, file_path: Path) -> str:
        """Extract content from text files"""
        encodings = ["utf-8", "utf-8-sig", "gb2312", "big5", "latin1"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        raise ValueError(f"Unable to decode file {file_path} with any encoding")

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        if not PDF_AVAILABLE:
            return "[PDF processing not available - install PyPDF2]"

        try:
            text_content = []
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1}: {e}")
                        text_content.append(
                            f"[Page {page_num + 1} - extraction failed]"
                        )

            return "\n\n".join(text_content)

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return f"[PDF extraction failed: {str(e)}]"

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            return "[DOCX processing not available - install python-docx]"

        try:
            doc = DocxDocument(file_path)
            paragraphs = [
                paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
            ]

            # Also extract from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)

            content = "\n".join(paragraphs)
            if tables_text:
                content += "\n\n[Tables]\n" + "\n".join(tables_text)

            return content

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return f"[DOCX extraction failed: {str(e)}]"

    def _extract_excel(self, file_path: Path) -> str:
        """Extract text from Excel files"""
        if not EXCEL_AVAILABLE:
            return "[Excel processing not available - install openpyxl]"

        try:
            workbook = load_workbook(file_path, data_only=True)
            sheets_content = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []

                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):  # Skip empty rows
                        sheet_data.append(" | ".join(row_data))

                if sheet_data:
                    sheets_content.append(
                        f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_data)
                    )

            return "\n\n".join(sheets_content)

        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return f"[Excel extraction failed: {str(e)}]"

    def _extract_csv(self, file_path: Path) -> str:
        """Extract text from CSV files"""
        try:
            import csv

            content_lines = []
            encodings = ["utf-8", "utf-8-sig", "gb2312", "big5"]

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding, newline="") as f:
                        csv_reader = csv.reader(f)
                        for row_num, row in enumerate(csv_reader):
                            if row and any(cell.strip() for cell in row):
                                content_lines.append(" | ".join(row))

                            # Limit rows to prevent huge files
                            if row_num > 1000:
                                content_lines.append("[... truncated after 1000 rows]")
                                break
                    break

                except UnicodeDecodeError:
                    continue

            return "\n".join(content_lines)

        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            return f"[CSV extraction failed: {str(e)}]"

    def _extract_json(self, file_path: Path) -> str:
        """Extract text from JSON files"""
        try:
            import json

            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            # Convert JSON to readable text format
            def json_to_text(obj, prefix=""):
                lines = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            lines.append(f"{prefix}{key}:")
                            lines.extend(json_to_text(value, prefix + "  "))
                        else:
                            lines.append(f"{prefix}{key}: {value}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        if isinstance(item, (dict, list)):
                            lines.append(f"{prefix}[{i}]:")
                            lines.extend(json_to_text(item, prefix + "  "))
                        else:
                            lines.append(f"{prefix}[{i}]: {item}")
                else:
                    lines.append(f"{prefix}{obj}")
                return lines

            return "\n".join(json_to_text(data))

        except Exception as e:
            logger.error(f"JSON extraction failed: {e}")
            return f"[JSON extraction failed: {str(e)}]"

    def _extract_html(self, file_path: Path) -> str:
        """Extract text from HTML files"""
        try:
            from html.parser import HTMLParser

            class HTMLTextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text_parts = []
                    self.ignore_tags = {"script", "style", "meta", "link"}
                    self.current_tag = None

                def handle_starttag(self, tag, attrs):
                    self.current_tag = tag

                def handle_data(self, data):
                    if self.current_tag not in self.ignore_tags:
                        text = data.strip()
                        if text:
                            self.text_parts.append(text)

                def get_text(self):
                    return "\n".join(self.text_parts)

            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()

            parser = HTMLTextExtractor()
            parser.feed(html_content)
            return parser.get_text()

        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return f"[HTML extraction failed: {str(e)}]"

    def _clean_content(self, content: str) -> str:
        """Clean and normalize extracted content"""
        if not content:
            return ""

        # Remove excessive whitespace
        content = re.sub(
            r"\n\s*\n\s*\n", "\n\n", content
        )  # Multiple newlines to double
        content = re.sub(
            r"[ \t]+", " ", content
        )  # Multiple spaces/tabs to single space
        content = re.sub(r"\r\n", "\n", content)  # Windows line endings

        # Remove common extraction artifacts
        content = re.sub(r"\x00", "", content)  # Null characters
        content = re.sub(
            r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", content
        )  # Control characters

        # Normalize Unicode
        try:
            content = content.encode("utf-8").decode("utf-8")
        except:
            pass

        return content.strip()

    def convert_traditional_to_simplified(self, text: str) -> str:
        """Convert Traditional Chinese to Simplified"""
        if self.cc_t2s and text:
            try:
                return self.cc_t2s.convert(text)
            except:
                logger.warning("Traditional to Simplified conversion failed")
        return text

    def convert_simplified_to_traditional(self, text: str) -> str:
        """Convert Simplified Chinese to Traditional"""
        if self.cc_s2t and text:
            try:
                return self.cc_s2t.convert(text)
            except:
                logger.warning("Simplified to Traditional conversion failed")
        return text

    def get_supported_formats(self) -> Dict[str, bool]:
        """Get supported file formats and their availability"""
        return {
            "text": True,
            "markdown": True,
            "csv": True,
            "json": True,
            "html": True,
            "pdf": PDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "excel": EXCEL_AVAILABLE,
        }
